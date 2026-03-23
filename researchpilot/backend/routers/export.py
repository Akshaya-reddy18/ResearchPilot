
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from db import get_session
from sqlalchemy.orm import Session
from models.user import Document, User
from utils.deps import get_current_user, require_workspace_member
import io

router = APIRouter(prefix="/api/export", tags=["export"])

# ================= PDF EXPORT =================

@router.get("/doc/{doc_id}/pdf")
def export_doc_pdf(
    doc_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user)
):
    if user is None:
        raise HTTPException(status_code=401, detail="Unauthorized")

    doc = session.query(Document).filter(Document.id == doc_id).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not require_workspace_member(doc.workspace_id, user, session):
        raise HTTPException(status_code=403, detail="Not allowed")

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from bs4 import BeautifulSoup
    from textwrap import wrap

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    soup = BeautifulSoup(doc.content or "", "html.parser")

    y = 750

    for tag in soup.find_all(["h1", "h2", "h3", "p"]):
        text = tag.get_text(strip=True)

        if tag.name == "h1":
            c.setFont("Helvetica-Bold", 18)
        elif tag.name == "h2":
            c.setFont("Helvetica-Bold", 16)
        elif tag.name == "h3":
            c.setFont("Helvetica-Bold", 14)
        else:
            c.setFont("Helvetica", 12)

        lines = wrap(text, 80)

        for line in lines:
            c.drawString(100, y, line)
            y -= 20

            if y < 50:
                c.showPage()
                y = 750

        y -= 10

    c.save()
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={doc.title}.pdf"
        }
    )


# ================= DOCX EXPORT =================

@router.get("/doc/{doc_id}/docx")
def export_doc_docx(
    doc_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user)
):
    if user is None:
        raise HTTPException(status_code=401, detail="Unauthorized")

    doc = session.query(Document).filter(Document.id == doc_id).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not require_workspace_member(doc.workspace_id, user, session):
        raise HTTPException(status_code=403, detail="Not allowed")

    from docx import Document as DocxDocument
    from bs4 import BeautifulSoup

    docx = DocxDocument()
    docx.add_heading(doc.title or "Document", 0)

    soup = BeautifulSoup(doc.content or "", "html.parser")

    for tag in soup.find_all(["h1", "h2", "h3", "p"]):
        text = tag.get_text(strip=True)

        if tag.name == "h1":
            docx.add_heading(text, level=1)
        elif tag.name == "h2":
            docx.add_heading(text, level=2)
        elif tag.name == "h3":
            docx.add_heading(text, level=3)
        else:
            docx.add_paragraph(text)

    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={doc.title}.docx"
        }
    )

